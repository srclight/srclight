"""DOCX extractor (requires python-docx)."""

from __future__ import annotations

import io
from pathlib import Path
from typing import TYPE_CHECKING

import docx

from .base import make_document, make_section

if TYPE_CHECKING:
    from ..db import Database


class DocxExtractor:
    language = "docx"
    extensions = (".docx",)

    def extract(
        self,
        file_id: int,
        rel_path: str,
        source: bytes,
        db: Database,
    ) -> int:
        stem = Path(rel_path).stem
        doc = docx.Document(io.BytesIO(source))

        # Detect headings
        headings = self._find_headings(doc)

        count = 0

        if headings:
            count += self._emit_heading_sections(
                file_id, rel_path, stem, doc, headings, db,
            )
        else:
            # No headings: whole doc as single symbol
            text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
            sym = make_document(file_id, rel_path, text, name=stem)
            db.insert_symbol(sym, rel_path)
            count = 1

        # Tables
        count += self._emit_tables(file_id, rel_path, stem, doc, db)

        return count

    def _find_headings(
        self, doc: docx.Document,
    ) -> list[tuple[int, str, int]]:
        """Find heading paragraphs.

        Returns list of (paragraph_index, heading_text, level).
        """
        headings: list[tuple[int, str, int]] = []
        for i, para in enumerate(doc.paragraphs):
            style_name = para.style.name if para.style else ""
            if style_name.startswith("Heading"):
                try:
                    level = int(style_name.split()[-1])
                except (ValueError, IndexError):
                    level = 1
                if para.text.strip():
                    headings.append((i, para.text.strip(), level))
        return headings

    def _emit_heading_sections(
        self,
        file_id: int,
        rel_path: str,
        stem: str,
        doc: docx.Document,
        headings: list[tuple[int, str, int]],
        db: Database,
    ) -> int:
        """Emit sections using heading-based hierarchy."""
        paragraphs = doc.paragraphs
        count = 0

        # Parent stack: list of (level, sym_id, name)
        parent_stack: list[tuple[int, int, str]] = []

        for h_idx, (para_idx, heading_text, level) in enumerate(headings):
            # Collect body paragraphs until next heading
            end_idx = headings[h_idx + 1][0] if h_idx + 1 < len(headings) else len(paragraphs)
            body_parts = []
            for j in range(para_idx + 1, end_idx):
                text = paragraphs[j].text.strip()
                if text:
                    body_parts.append(text)
            body = "\n\n".join(body_parts)

            content = f"{heading_text}\n\n{body}".strip() if body else heading_text

            # Maintain parent stack by level
            while parent_stack and parent_stack[-1][0] >= level:
                parent_stack.pop()

            parent_id = parent_stack[-1][1] if parent_stack else None

            # Build qualified name from stack
            ancestry = [name for _, _, name in parent_stack] + [heading_text]
            qualified = f"{stem} > " + " > ".join(ancestry)

            # First body paragraph as doc_comment for porter-stemmed FTS3
            doc_comment = body_parts[0][:200] if body_parts else None

            sym = make_section(
                file_id, rel_path,
                name=heading_text,
                qualified_name=qualified,
                content=content,
                start_line=para_idx + 1,
                end_line=end_idx,
                signature=f"{'#' * level} {heading_text}",
                parent_symbol_id=parent_id,
                doc_comment=doc_comment,
            )
            sym_id = db.insert_symbol(sym, rel_path)
            parent_stack.append((level, sym_id, heading_text))
            count += 1

        return count

    def _emit_tables(
        self,
        file_id: int,
        rel_path: str,
        stem: str,
        doc: docx.Document,
        db: Database,
    ) -> int:
        """Emit table symbols."""
        count = 0
        for t_idx, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append("\t".join(cells))
            content = "\n".join(rows)
            if not content.strip():
                continue

            name = f"Table {t_idx + 1}"
            sym = make_section(
                file_id, rel_path,
                name=name,
                qualified_name=f"{stem} > {name}",
                content=content,
                kind="table",
            )
            db.insert_symbol(sym, rel_path)
            count += 1

        return count
