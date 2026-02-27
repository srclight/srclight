"""Tests for document extractors."""

import csv
import io
from pathlib import Path
from unittest.mock import patch

import pytest

from srclight.db import Database
from srclight.indexer import IndexConfig, Indexer


def _paddle_available() -> bool:
    """Check if paddleocr and pdf2image are installed."""
    try:
        import paddleocr  # noqa: F401
        import pdf2image  # noqa: F401
        return True
    except ImportError:
        return False


@pytest.fixture
def db(tmp_path):
    db_path = tmp_path / "test.db"
    db = Database(db_path)
    db.open()
    db.initialize()
    yield db
    db.close()


# ---------------------------------------------------------------------------
# Stdlib: Text extractor
# ---------------------------------------------------------------------------


class TestTextExtractor:
    def test_plain_text_small(self, db, tmp_path):
        """Small plain text → single document symbol."""
        project = tmp_path / "proj"
        project.mkdir()
        (project / "notes.txt").write_text("Hello world.\nThis is a note.")

        indexer = Indexer(db, IndexConfig(root=project))
        stats = indexer.index(project)
        assert stats.files_indexed == 1
        assert stats.symbols_extracted == 1

        syms = db.search_symbols("notes")
        assert len(syms) >= 1
        assert syms[0]["kind"] == "document"

    def test_plain_text_large(self, db, tmp_path):
        """Large plain text → chunked sections."""
        project = tmp_path / "proj"
        project.mkdir()
        lines = [f"Line {i}" for i in range(200)]
        (project / "big.txt").write_text("\n".join(lines))

        indexer = Indexer(db, IndexConfig(root=project))
        stats = indexer.index(project)
        assert stats.files_indexed == 1
        # 200 lines / 80-line chunks → 3 sections
        assert stats.symbols_extracted == 3

        syms = db.search_symbols("Line 0")
        assert len(syms) >= 1
        assert syms[0]["kind"] == "section"

    def test_rst_headings(self, db, tmp_path):
        """RST with underlined headings → sections."""
        project = tmp_path / "proj"
        project.mkdir()
        rst = """\
Introduction
============

This is the intro paragraph.

Installation
------------

Run pip install.

Usage
-----

Import and use.
"""
        (project / "guide.rst").write_text(rst)

        indexer = Indexer(db, IndexConfig(root=project))
        stats = indexer.index(project)
        assert stats.files_indexed == 1
        assert stats.symbols_extracted == 3  # Introduction, Installation, Usage

        syms = db.search_symbols("Introduction")
        assert len(syms) >= 1
        assert syms[0]["kind"] == "section"


# ---------------------------------------------------------------------------
# Stdlib: Email extractor
# ---------------------------------------------------------------------------


class TestEmailExtractor:
    def test_simple_email(self, db, tmp_path):
        """Parse a simple .eml file."""
        project = tmp_path / "proj"
        project.mkdir()

        eml = """\
From: alice@example.com
To: bob@example.com
Subject: Meeting Tomorrow
Date: Mon, 1 Jan 2024 10:00:00 +0000
Content-Type: text/plain; charset="utf-8"

Hi Bob,

Can we meet tomorrow at 3pm?

Thanks,
Alice
"""
        (project / "meeting.eml").write_bytes(eml.encode("utf-8"))

        indexer = Indexer(db, IndexConfig(root=project))
        stats = indexer.index(project)
        assert stats.files_indexed == 1
        assert stats.symbols_extracted == 1

        syms = db.search_symbols("Meeting Tomorrow")
        assert len(syms) >= 1
        sym = syms[0]
        assert sym["kind"] == "document"
        # Check doc_comment via full symbol record
        full = db.get_symbol_by_id(sym["symbol_id"])
        assert full is not None
        assert "alice@example.com" in (full.doc_comment or "")


# ---------------------------------------------------------------------------
# Stdlib: CSV extractor
# ---------------------------------------------------------------------------


class TestCsvExtractor:
    def test_simple_csv(self, db, tmp_path):
        """Parse a simple CSV file."""
        project = tmp_path / "proj"
        project.mkdir()

        buf = io.StringIO()
        writer = csv.writer(buf)
        writer.writerow(["Name", "Age", "City"])
        writer.writerow(["Alice", "30", "NYC"])
        writer.writerow(["Bob", "25", "LA"])
        (project / "people.csv").write_text(buf.getvalue())

        indexer = Indexer(db, IndexConfig(root=project))
        stats = indexer.index(project)
        assert stats.files_indexed == 1
        assert stats.symbols_extracted == 1

        syms = db.search_symbols("people")
        assert len(syms) >= 1
        sym = syms[0]
        assert sym["kind"] == "document"
        full = db.get_symbol_by_id(sym["symbol_id"])
        assert full is not None
        assert "Name" in (full.signature or "")
        assert "Columns:" in (full.doc_comment or "")
        assert "3 rows" in (full.doc_comment or "")

    def test_tsv(self, db, tmp_path):
        """Parse a TSV file."""
        project = tmp_path / "proj"
        project.mkdir()
        (project / "data.tsv").write_text("col1\tcol2\nval1\tval2\n")

        indexer = Indexer(db, IndexConfig(root=project))
        stats = indexer.index(project)
        assert stats.files_indexed == 1
        assert stats.symbols_extracted == 1


# ---------------------------------------------------------------------------
# Optional: PDF extractor
# ---------------------------------------------------------------------------


class TestPdfExtractor:
    @pytest.fixture(autouse=True)
    def _skip_if_no_dep(self):
        pytest.importorskip("pdfplumber")

    def test_pdf_indexing(self, db, tmp_path):
        """Create a simple PDF and index it."""
        # We need reportlab or fpdf to create a test PDF, but we can test
        # with a minimal valid PDF binary.
        fpdf2 = pytest.importorskip("fpdf")

        project = tmp_path / "proj"
        project.mkdir()

        pdf = fpdf2.FPDF()
        pdf.add_page()
        pdf.set_font("Helvetica", size=12)
        pdf.cell(text="Hello PDF World")
        pdf.output(str(project / "test.pdf"))

        indexer = Indexer(db, IndexConfig(root=project))
        stats = indexer.index(project)
        assert stats.files_indexed == 1
        assert stats.symbols_extracted >= 1


# ---------------------------------------------------------------------------
# Optional: PDF OCR fallback (PaddleOCR)
# ---------------------------------------------------------------------------


class TestPdfOcr:
    @pytest.fixture(autouse=True)
    def _skip_if_no_dep(self):
        pytest.importorskip("pdfplumber")

    def test_ocr_fallback_wired(self, db, tmp_path):
        """When extract_text returns nothing but page has images, OCR path
        is attempted. If PaddleOCR unavailable, gracefully falls back to
        empty document (no crash)."""
        fpdf2 = pytest.importorskip("fpdf")

        project = tmp_path / "proj"
        project.mkdir()

        # Create a normal PDF (has native text)
        pdf = fpdf2.FPDF()
        pdf.add_page()
        pdf.set_font("Helvetica", size=12)
        pdf.cell(text="Hello OCR test")
        pdf.output(str(project / "scan.pdf"))

        # Patch extract_text to return "" (simulating scanned page)
        # and page.images to be non-empty (simulating embedded images)
        with patch("pdfplumber.page.Page.extract_text", return_value=""), \
             patch("pdfplumber.page.Page.extract_words", return_value=[]), \
             patch.object(
                 type(next(iter([None]))),  # dummy — we'll patch the property below
                 "images", new_callable=lambda: property(lambda self: []),
             ) if False else \
             patch(
                 "srclight.extractors.pdf_extractor.PdfExtractor._init_paddle",
                 return_value=False,
             ):
            indexer = Indexer(db, IndexConfig(root=project))
            stats = indexer.index(project)
            # Should produce at least the empty-document fallback symbol
            assert stats.files_indexed == 1
            assert stats.symbols_extracted >= 1

    def test_native_pdf_skips_ocr(self, db, tmp_path):
        """PDFs with native text should never trigger OCR."""
        fpdf2 = pytest.importorskip("fpdf")

        project = tmp_path / "proj"
        project.mkdir()

        pdf = fpdf2.FPDF()
        pdf.add_page()
        pdf.set_font("Helvetica", size=12)
        pdf.cell(text="Native text page")
        pdf.output(str(project / "native.pdf"))

        with patch(
            "srclight.extractors.pdf_extractor.PdfExtractor._ocr_page",
            wraps=None,
        ) as mock_ocr:
            mock_ocr.return_value = None
            indexer = Indexer(db, IndexConfig(root=project))
            stats = indexer.index(project)
            assert stats.files_indexed == 1
            assert stats.symbols_extracted >= 1
            # _ocr_page should never have been called since native text exists
            mock_ocr.assert_not_called()

    @pytest.mark.skipif(
        not _paddle_available(),
        reason="paddleocr or pdf2image not installed",
    )
    def test_paddle_ocr_integration(self, db, tmp_path):
        """End-to-end: create an image-only PDF and verify OCR extracts text."""
        fpdf2 = pytest.importorskip("fpdf")
        from PIL import Image, ImageDraw, ImageFont

        project = tmp_path / "proj"
        project.mkdir()

        # Create an image with text
        img = Image.new("RGB", (400, 100), color="white")
        draw = ImageDraw.Draw(img)
        draw.text((10, 30), "Scanned document text", fill="black")
        img_path = tmp_path / "page.png"
        img.save(str(img_path))

        # Create a PDF with only that image (no native text)
        pdf = fpdf2.FPDF()
        pdf.add_page()
        pdf.image(str(img_path), x=10, y=10, w=180)
        pdf.output(str(project / "scanned.pdf"))

        indexer = Indexer(db, IndexConfig(root=project))
        stats = indexer.index(project)
        assert stats.files_indexed == 1
        assert stats.symbols_extracted >= 1

        # Verify some OCR'd text made it into a symbol
        syms = db.search_symbols("Scanned")
        assert len(syms) >= 1


# ---------------------------------------------------------------------------
# Optional: DOCX extractor
# ---------------------------------------------------------------------------


class TestDocxExtractor:
    @pytest.fixture(autouse=True)
    def _skip_if_no_dep(self):
        pytest.importorskip("docx")

    def test_docx_with_headings(self, db, tmp_path):
        """Create a DOCX with headings and index it."""
        import docx

        project = tmp_path / "proj"
        project.mkdir()

        doc = docx.Document()
        doc.add_heading("Chapter 1", level=1)
        doc.add_paragraph("This is the first chapter.")
        doc.add_heading("Section 1.1", level=2)
        doc.add_paragraph("Subsection content.")
        doc.add_heading("Chapter 2", level=1)
        doc.add_paragraph("Second chapter content.")
        doc.save(str(project / "report.docx"))

        indexer = Indexer(db, IndexConfig(root=project))
        stats = indexer.index(project)
        assert stats.files_indexed == 1
        # 3 heading sections: Chapter 1, Section 1.1, Chapter 2
        assert stats.symbols_extracted == 3

        syms = db.search_symbols("Chapter 1")
        assert len(syms) >= 1
        assert syms[0]["kind"] == "section"

    def test_docx_no_headings(self, db, tmp_path):
        """DOCX without headings → single document."""
        import docx

        project = tmp_path / "proj"
        project.mkdir()

        doc = docx.Document()
        doc.add_paragraph("Just a plain document.")
        doc.save(str(project / "plain.docx"))

        indexer = Indexer(db, IndexConfig(root=project))
        stats = indexer.index(project)
        assert stats.files_indexed == 1
        assert stats.symbols_extracted == 1

        syms = db.search_symbols("plain")
        assert len(syms) >= 1
        assert syms[0]["kind"] == "document"


# ---------------------------------------------------------------------------
# Optional: XLSX extractor
# ---------------------------------------------------------------------------


class TestXlsxExtractor:
    @pytest.fixture(autouse=True)
    def _skip_if_no_dep(self):
        pytest.importorskip("openpyxl")

    def test_xlsx_single_sheet(self, db, tmp_path):
        """Index a simple XLSX with one sheet."""
        import openpyxl

        project = tmp_path / "proj"
        project.mkdir()

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Data"
        ws.append(["Name", "Value"])
        ws.append(["Alpha", 1])
        ws.append(["Beta", 2])
        wb.save(str(project / "data.xlsx"))

        indexer = Indexer(db, IndexConfig(root=project))
        stats = indexer.index(project)
        assert stats.files_indexed == 1
        assert stats.symbols_extracted == 1  # One sheet = one section

        syms = db.search_symbols("Data")
        assert len(syms) >= 1
        assert syms[0]["kind"] == "section"
        full = db.get_symbol_by_id(syms[0]["symbol_id"])
        assert full is not None
        assert "Name" in (full.signature or "")

    def test_xlsx_multiple_sheets(self, db, tmp_path):
        """Index XLSX with multiple sheets."""
        import openpyxl

        project = tmp_path / "proj"
        project.mkdir()

        wb = openpyxl.Workbook()
        ws1 = wb.active
        ws1.title = "Revenue"
        ws1.append(["Month", "Amount"])
        ws1.append(["Jan", 100])

        ws2 = wb.create_sheet("Costs")
        ws2.append(["Category", "Amount"])
        ws2.append(["Rent", 500])
        wb.save(str(project / "finance.xlsx"))

        indexer = Indexer(db, IndexConfig(root=project))
        stats = indexer.index(project)
        assert stats.files_indexed == 1
        assert stats.symbols_extracted == 2  # Two sheets


# ---------------------------------------------------------------------------
# Optional: HTML extractor
# ---------------------------------------------------------------------------


class TestHtmlExtractor:
    @pytest.fixture(autouse=True)
    def _skip_if_no_dep(self):
        pytest.importorskip("bs4")

    def test_html_with_headings(self, db, tmp_path):
        """HTML with headings → sections."""
        project = tmp_path / "proj"
        project.mkdir()

        html = """\
<html>
<head><title>My Guide</title></head>
<body>
<h1>Introduction</h1>
<p>Welcome to the guide.</p>
<h2>Getting Started</h2>
<p>First steps here.</p>
<h1>Advanced</h1>
<p>More details.</p>
</body>
</html>
"""
        (project / "guide.html").write_text(html)

        indexer = Indexer(db, IndexConfig(root=project))
        stats = indexer.index(project)
        assert stats.files_indexed == 1
        assert stats.symbols_extracted == 3  # Introduction, Getting Started, Advanced

        syms = db.search_symbols("Introduction")
        assert len(syms) >= 1
        assert syms[0]["kind"] == "section"

    def test_html_no_headings(self, db, tmp_path):
        """HTML without headings → single document."""
        project = tmp_path / "proj"
        project.mkdir()

        html = "<html><body><p>Just a paragraph.</p></body></html>"
        (project / "simple.html").write_text(html)

        indexer = Indexer(db, IndexConfig(root=project))
        stats = indexer.index(project)
        assert stats.files_indexed == 1
        assert stats.symbols_extracted == 1


# ---------------------------------------------------------------------------
# Optional: Image extractor
# ---------------------------------------------------------------------------


class TestImageExtractor:
    @pytest.fixture(autouse=True)
    def _skip_if_no_dep(self):
        pytest.importorskip("PIL")

    def test_png_metadata(self, db, tmp_path):
        """Index a small PNG for metadata."""
        from PIL import Image

        project = tmp_path / "proj"
        project.mkdir()

        img = Image.new("RGB", (100, 50), color="red")
        img.save(str(project / "red.png"))

        indexer = Indexer(db, IndexConfig(root=project))
        stats = indexer.index(project)
        assert stats.files_indexed == 1
        assert stats.symbols_extracted == 1

        syms = db.search_symbols("red")
        assert len(syms) >= 1
        full = db.get_symbol_by_id(syms[0]["symbol_id"])
        assert full is not None
        assert "100x50" in (full.signature or "")

    def test_svg_extraction(self, db, tmp_path):
        """Index an SVG file."""
        project = tmp_path / "proj"
        project.mkdir()

        svg = """\
<svg xmlns="http://www.w3.org/2000/svg" width="200" height="100" viewBox="0 0 200 100">
  <title>Test Icon</title>
  <desc>A simple test icon</desc>
  <rect width="200" height="100" fill="blue"/>
</svg>
"""
        (project / "icon.svg").write_text(svg)

        indexer = Indexer(db, IndexConfig(root=project))
        stats = indexer.index(project)
        assert stats.files_indexed == 1
        assert stats.symbols_extracted == 1

        syms = db.search_symbols("icon")
        assert len(syms) >= 1
        full = db.get_symbol_by_id(syms[0]["symbol_id"])
        assert full is not None
        assert "SVG" in (full.signature or "")


# ---------------------------------------------------------------------------
# Search quality: targeted tests for FTS/embedding improvements
# ---------------------------------------------------------------------------


class TestSearchQualityPdf:
    """Tests for PDF search quality fixes."""

    @pytest.fixture(autouse=True)
    def _skip_if_no_dep(self):
        pytest.importorskip("pdfplumber")

    def test_multi_word_heading_not_shattered(self, db, tmp_path):
        """Multi-word headings must be kept as a single string."""
        fpdf2 = pytest.importorskip("fpdf")

        project = tmp_path / "proj"
        project.mkdir()

        pdf = fpdf2.FPDF()
        pdf.add_page()
        # Large heading text (multiple words at heading size)
        pdf.set_font("Helvetica", size=24)
        pdf.cell(text="Chapter One Introduction")
        pdf.ln()
        # Body text at smaller size
        pdf.set_font("Helvetica", size=12)
        pdf.cell(text="This is the body of chapter one.")
        pdf.ln()
        # Second heading
        pdf.set_font("Helvetica", size=24)
        pdf.cell(text="Chapter Two Methods")
        pdf.ln()
        pdf.set_font("Helvetica", size=12)
        pdf.cell(text="Methods description here.")
        pdf.output(str(project / "report.pdf"))

        indexer = Indexer(db, IndexConfig(root=project))
        stats = indexer.index(project)

        # Should be 2 sections (not 5+ shattered words)
        assert stats.symbols_extracted == 2

        syms = db.search_symbols("Chapter One Introduction")
        assert len(syms) >= 1
        assert syms[0]["name"] == "Chapter One Introduction"

    def test_pdf_sections_have_doc_comment(self, db, tmp_path):
        """Each PDF section should get a doc_comment for porter-stemmed search."""
        fpdf2 = pytest.importorskip("fpdf")

        project = tmp_path / "proj"
        project.mkdir()

        pdf = fpdf2.FPDF()
        pdf.add_page()
        pdf.set_font("Helvetica", size=20)
        pdf.cell(text="First Heading")
        pdf.ln()
        pdf.set_font("Helvetica", size=12)
        pdf.cell(text="First body paragraph content.")
        pdf.ln()
        pdf.set_font("Helvetica", size=20)
        pdf.cell(text="Second Heading")
        pdf.ln()
        pdf.set_font("Helvetica", size=12)
        pdf.cell(text="Second body paragraph content.")
        pdf.output(str(project / "doc.pdf"))

        indexer = Indexer(db, IndexConfig(root=project))
        indexer.index(project)

        syms = db.search_symbols("Second Heading")
        assert len(syms) >= 1
        full = db.get_symbol_by_id(syms[0]["symbol_id"])
        assert full is not None
        # Second section should have doc_comment (not None)
        assert full.doc_comment is not None
        assert len(full.doc_comment) > 0


class TestSearchQualityDocx:
    """Tests for DOCX search quality fixes."""

    @pytest.fixture(autouse=True)
    def _skip_if_no_dep(self):
        pytest.importorskip("docx")

    def test_docx_sections_have_doc_comment(self, db, tmp_path):
        """DOCX sections should have doc_comment populated."""
        import docx

        project = tmp_path / "proj"
        project.mkdir()

        doc = docx.Document()
        doc.add_heading("Authentication", level=1)
        doc.add_paragraph("Users authenticate with OAuth tokens.")
        doc.add_heading("Authorization", level=1)
        doc.add_paragraph("Role-based access control applies.")
        doc.save(str(project / "security.docx"))

        indexer = Indexer(db, IndexConfig(root=project))
        indexer.index(project)

        syms = db.search_symbols("Authentication")
        assert len(syms) >= 1
        full = db.get_symbol_by_id(syms[0]["symbol_id"])
        assert full is not None
        assert full.doc_comment is not None
        assert "OAuth" in full.doc_comment


class TestSearchQualityText:
    """Tests for text chunk naming."""

    def test_chunk_names_from_content(self, db, tmp_path):
        """Text chunks should use first line as name, not 'Section N'."""
        project = tmp_path / "proj"
        project.mkdir()

        lines = [f"Configuration settings for module {i}" for i in range(200)]
        (project / "config.txt").write_text("\n".join(lines))

        indexer = Indexer(db, IndexConfig(root=project))
        indexer.index(project)

        syms = db.search_symbols("Configuration settings for module 0")
        assert len(syms) >= 1
        # Name should be derived from content, not generic "Section 1"
        assert "Section 1" not in syms[0]["name"]
        assert "Configuration" in syms[0]["name"]


class TestSearchQualityEmail:
    """Tests for email HTML stripping."""

    def test_html_only_email_stripped(self, db, tmp_path):
        """HTML-only emails should have tags stripped from content."""
        project = tmp_path / "proj"
        project.mkdir()

        eml = """\
From: sender@example.com
To: recipient@example.com
Subject: HTML Newsletter
Content-Type: text/html; charset="utf-8"

<html><body><div><h1>Welcome</h1><p>Hello subscriber!</p></div></body></html>
"""
        (project / "newsletter.eml").write_bytes(eml.encode("utf-8"))

        indexer = Indexer(db, IndexConfig(root=project))
        indexer.index(project)

        syms = db.search_symbols("HTML Newsletter")
        assert len(syms) >= 1
        full = db.get_symbol_by_id(syms[0]["symbol_id"])
        assert full is not None
        # Content should not contain HTML tags
        assert "<div>" not in full.content
        assert "<h1>" not in full.content
        assert "Welcome" in full.content
        assert "Hello subscriber" in full.content


class TestSearchQualityCsv:
    """Tests for CSV/XLSX column names in doc_comment."""

    def test_csv_column_names_in_doc_comment(self, db, tmp_path):
        """CSV doc_comment should contain column names."""
        project = tmp_path / "proj"
        project.mkdir()

        buf = io.StringIO()
        writer = csv.writer(buf)
        writer.writerow(["customer_id", "order_date", "total_amount"])
        writer.writerow(["1", "2024-01-01", "99.99"])
        (project / "orders.csv").write_text(buf.getvalue())

        indexer = Indexer(db, IndexConfig(root=project))
        indexer.index(project)

        syms = db.search_symbols("orders")
        assert len(syms) >= 1
        full = db.get_symbol_by_id(syms[0]["symbol_id"])
        assert full is not None
        assert "customer_id" in full.doc_comment
        assert "order_date" in full.doc_comment
        assert "Columns:" in full.doc_comment


class TestSearchQualityXlsx:
    """Tests for XLSX column names in doc_comment."""

    @pytest.fixture(autouse=True)
    def _skip_if_no_dep(self):
        pytest.importorskip("openpyxl")

    def test_xlsx_column_names_in_doc_comment(self, db, tmp_path):
        """XLSX doc_comment should contain column names."""
        import openpyxl

        project = tmp_path / "proj"
        project.mkdir()

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Sales"
        ws.append(["product_name", "quantity", "unit_price"])
        ws.append(["Widget", 10, 5.99])
        wb.save(str(project / "sales.xlsx"))

        indexer = Indexer(db, IndexConfig(root=project))
        indexer.index(project)

        syms = db.search_symbols("Sales")
        assert len(syms) >= 1
        full = db.get_symbol_by_id(syms[0]["symbol_id"])
        assert full is not None
        assert "product_name" in full.doc_comment
        assert "quantity" in full.doc_comment
        assert "Columns:" in full.doc_comment


class TestSearchQualityHtml:
    """Tests for HTML doc_comment on all sections."""

    @pytest.fixture(autouse=True)
    def _skip_if_no_dep(self):
        pytest.importorskip("bs4")

    def test_html_all_sections_have_doc_comment(self, db, tmp_path):
        """All HTML sections should have doc_comment, not just the first."""
        project = tmp_path / "proj"
        project.mkdir()

        html = """\
<html>
<head><title>API Docs</title></head>
<body>
<h1>Overview</h1>
<p>API overview paragraph.</p>
<h1>Authentication</h1>
<p>Use bearer tokens for authentication.</p>
<h1>Endpoints</h1>
<p>List of available endpoints.</p>
</body>
</html>
"""
        (project / "api.html").write_text(html)

        indexer = Indexer(db, IndexConfig(root=project))
        indexer.index(project)

        # Check that the second and third sections have doc_comment
        syms = db.search_symbols("Authentication")
        assert len(syms) >= 1
        full = db.get_symbol_by_id(syms[0]["symbol_id"])
        assert full is not None
        assert full.doc_comment is not None
        assert "bearer" in full.doc_comment.lower() or "token" in full.doc_comment.lower()


# ---------------------------------------------------------------------------
# Integration: mixed content directory
# ---------------------------------------------------------------------------


class TestMixedContent:
    def test_code_and_text(self, db, tmp_path):
        """Index a directory with both code and text files."""
        project = tmp_path / "proj"
        project.mkdir()

        (project / "main.py").write_text("def hello(): pass\n")
        (project / "README.md").write_text("# Hello\n\nA project.\n")
        (project / "notes.txt").write_text("Some notes here.")

        indexer = Indexer(db, IndexConfig(root=project))
        stats = indexer.index(project)
        assert stats.files_indexed == 3
        # 1 Python function + 1 markdown section + 1 text document
        assert stats.symbols_extracted == 3

    def test_csv_and_code(self, db, tmp_path):
        """CSV alongside code files."""
        project = tmp_path / "proj"
        project.mkdir()

        (project / "data.csv").write_text("a,b,c\n1,2,3\n")
        (project / "process.py").write_text("def process(): pass\n")

        indexer = Indexer(db, IndexConfig(root=project))
        stats = indexer.index(project)
        assert stats.files_indexed == 2
        assert stats.symbols_extracted == 2


# ---------------------------------------------------------------------------
# Registry tests
# ---------------------------------------------------------------------------


class TestExtractorRegistry:
    def test_stdlib_extractors_always_registered(self):
        """Stdlib extractors are always available."""
        from srclight.extractors import DOCUMENT_EXTENSIONS, get_registry

        registry = get_registry()
        assert "text" in registry
        assert "email" in registry
        assert "csv" in registry
        assert ".txt" in DOCUMENT_EXTENSIONS
        assert ".eml" in DOCUMENT_EXTENSIONS
        assert ".csv" in DOCUMENT_EXTENSIONS

    def test_detect_document_language(self):
        from srclight.extractors import detect_document_language

        assert detect_document_language(".txt") == "text"
        assert detect_document_language(".TXT") == "text"
        assert detect_document_language(".eml") == "email"
        assert detect_document_language(".csv") == "csv"
        assert detect_document_language(".py") is None
        assert detect_document_language(".unknown") is None
